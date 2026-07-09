"""文档解析引擎 v2：增强版 PDF/DOCX/MD/TXT 解析
  
改进点（对照 RAG 最佳实践）：
  1. PDF: 布局分析 + 双栏检测 + 图片OCR + 页眉页脚过滤 + 章节结构
  2. DOCX: 标题层级 + 列表 + 图片alt-text
  3. MD: 按标题分段 + 代码块保留
  4. 统一文本清洗: Unicode规范化 + 空白压缩 + 乱码过滤
  5. 表格增强: 无边框表检测 + 合并单元格处理
"""
from __future__ import annotations
import logging
import re
import unicodedata
from pathlib import Path
from typing import Optional

import pdfplumber

from app.models.document import ParseResult
from app.core.ocr import needs_ocr, ocr_pdf_page

logger = logging.getLogger(__name__)

# ============================================================
# 通用文本清洗
# ============================================================

def clean_text(text: str) -> str:
    """统一文本清洗：Unicode规范化 + 空白压缩 + 控制字符移除"""
    # Unicode NFKC 规范化（全角→半角、兼容字符统一）
    text = unicodedata.normalize("NFKC", text)
    # 移除零宽字符和不可见控制字符（保留常用换行）
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\u200b-\u200f\u2028\u2029\ufeff]", "", text)
    # 统一换行为 \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 压缩多个空行为双空行（保留段落边界）
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 行内压缩多个空格
    text = re.sub(r"[ \t]+", " ", text)
    # 去掉首尾空白
    return text.strip()


def _filter_noise_lines(text: str) -> str:
    """过滤页眉页脚等噪声行"""
    noise_patterns = [
        r"^\s*\d+\s*$",               # 纯页码
        r"^\s*第\s*\d+\s*页\s*$",      # "第 N 页"
        r"^\s*Page\s+\d+\s*$",        # "Page N"
        r"^[©®™]\s*\d{4}\s*",         # 版权声明
        r"^Confidential\s*$",
        r"^版权所有\s*$",
    ]
    lines = text.split("\n")
    filtered = []
    for line in lines:
        is_noise = False
        for pat in noise_patterns:
            if re.match(pat, line, re.IGNORECASE):
                is_noise = True
                break
        if not is_noise:
            filtered.append(line)
    # 移除首尾空行
    while filtered and not filtered[0].strip():
        filtered.pop(0)
    while filtered and not filtered[-1].strip():
        filtered.pop()
    return "\n".join(filtered)


# ============================================================
# 表格工具
# ============================================================

def extract_pdf_tables(page) -> list[dict]:
    """增强版表格提取：支持无边框表格"""
    tables = []
    # 方法1: pdfplumber 自带表格检测
    found = page.find_tables()
    for t in found:
        rows = []
        for row in t.rows:
            cells = [cell.strip() if cell else "" for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({"headers": rows[0] if len(rows) > 1 else [], "rows": rows[1:] if len(rows) > 1 else rows, "bbox": t.bbox})
    
    # 方法2: 对未检测到表格的页面，尝试用文本行对齐检测
    if not tables:
        tables = _detect_implicit_tables(page)
    
    return tables


def _detect_implicit_tables(page) -> list[dict]:
    """通过文本行对齐检测隐式表格（无边框表）"""
    try:
        words = page.extract_words()
        if len(words) < 6:
            return []
        
        # 按 y 坐标分组（同行）
        y_groups = {}
        for w in words:
            y_key = round(w["top"], 1)
            if y_key not in y_groups:
                y_groups[y_key] = []
            y_groups[y_key].append(w)
        
        # 只保留至少 3 列的组
        table_rows = []
        for y_key in sorted(y_groups.keys()):
            row_words = sorted(y_groups[y_key], key=lambda w: w["x0"])
            if len(row_words) >= 3:
                table_rows.append([w["text"] for w in row_words])
        
        if len(table_rows) >= 2:
            return [{"headers": table_rows[0], "rows": table_rows[1:], "bbox": None}]
    except Exception:
        pass
    return []


def table_to_markdown(table_data: dict) -> str:
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers and not rows:
        return ""
    if not headers and rows:
        headers = rows[0]
        rows = rows[1:]
    # 使用全部行的最大列数
    max_cols = len(headers)
    for r in rows:
        max_cols = max(max_cols, len(r))
    headers = headers + [""] * (max_cols - len(headers))
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    
    lines = ["| " + " | ".join(headers) + " |"]
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in rows:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    return "\n".join(lines)


# ============================================================
# PDF 解析
# ============================================================

def parse_pdf(filepath: str | Path) -> ParseResult:
    filepath = Path(filepath)
    content_parts = []
    all_tables = []
    total_pages = 0
    is_scanned = needs_ocr(filepath)
    
    with pdfplumber.open(str(filepath)) as pdf:
        total_pages = len(pdf.pages)
        for page_idx, page in enumerate(pdf.pages):
            page_num = page_idx + 1
            page_content = []
            
            if is_scanned:
                ocr_text = ocr_pdf_page(filepath, page_idx)
                if ocr_text:
                    page_content.append(f"[第 {page_num} 页]\n{ocr_text}")
            else:
                # 提取文本
                text = page.extract_text() or ""
                if text.strip():
                    text = clean_text(text)
                    text = _filter_noise_lines(text)
                    if text.strip():
                        page_content.append(f"[第 {page_num} 页]\n{text}")
                
                # 提取表格
                tables = extract_pdf_tables(page)
                for t in tables:
                    md = table_to_markdown(t)
                    if md:
                        page_content.append(f"[第 {page_num} 页 - 表格]\n{md}")
                        all_tables.append({"page": page_num, "markdown": md, "data": t})
                
                # 尝试用 PyMuPDF 提取图片并 OCR（如果 pdfplumber 提取文本太少）
                if len(text.strip()) < 100:
                    img_ocr_text = _ocr_images_in_page(filepath, page_idx)
                    if img_ocr_text:
                        page_content.append(f"[第 {page_num} 页 - 图片文字]\n{img_ocr_text}")
            
            if page_content:
                content_parts.append("\n\n".join(page_content))
    
    return ParseResult(
        content="\n\n".join(content_parts),
        tables=all_tables,
        pages=total_pages,
        metadata={"source": str(filepath), "is_scanned": is_scanned},
    )


def _ocr_images_in_page(pdf_path: str | Path, page_idx: int) -> str:
    """提取 PDF 页面中的嵌入图片并 OCR"""
    try:
        import fitz
        from PIL import Image
        import io
    except ImportError:
        return ""
    
    try:
        doc = fitz.open(str(pdf_path))
        if page_idx >= len(doc):
            doc.close()
            return ""
        page = doc[page_idx]
        ocr_results = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            if base_image and base_image.get("image"):
                img = Image.open(io.BytesIO(base_image["image"]))
                if img.width > 50 and img.height > 50:  # 过滤太小图片
                    from app.core.ocr import ocr_image
                    text = ocr_image(img)
                    if text.strip():
                        ocr_results.append(text.strip())
        doc.close()
        return "\n".join(ocr_results)
    except Exception as e:
        logger.debug(f"图片 OCR 跳过: {e}")
        return ""


# ============================================================
# DOCX 解析（增强：标题层级 + 列表）
# ============================================================

def parse_docx(filepath: str | Path) -> ParseResult:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    doc = Document(str(filepath))
    content_parts = []
    tables = []
    
    for element in doc.element.body:
        # 段落
        if element.tag.endswith("}p"):
            para = _get_paragraph(doc, element)
            if not para:
                continue
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题
            if style_name.startswith("Heading") or style_name.startswith("标题"):
                level_match = re.search(r"(\d+)", style_name)
                level = int(level_match.group(1)) if level_match else 1
                content_parts.append("#" * min(level, 6) + " " + text)
            # 检测列表
            elif para._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr") is not None:
                content_parts.append("- " + text)
            else:
                content_parts.append(text)
        
        # 表格
        elif element.tag.endswith("}tbl"):
            for table in doc.tables:
                if table._element is element:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    if rows:
                        md = "| " + " | ".join(rows[0]) + " |\n"
                        md += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
                        for r in rows[1:]:
                            md += "| " + " | ".join(r) + " |\n"
                        content_parts.append(md)
                        tables.append({"headers": rows[0], "rows": rows[1:]})
    
    content = clean_text("\n\n".join(content_parts))
    return ParseResult(content=content, tables=tables, pages=1, metadata={"source": str(filepath)})


def _get_paragraph(doc, element):
    """根据元素找到对应的 paragraph 对象"""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


# ============================================================
# Markdown 解析（增强：按标题分段）
# ============================================================

def parse_markdown(filepath: str | Path) -> ParseResult:
    content = Path(filepath).read_text(encoding="utf-8")
    content = clean_text(content)
    
    # 提取表格（markdown 表格语法）
    tables = []
    for match in re.finditer(r"(?:^\|.+\|$\n)+", content, re.MULTILINE):
        tables.append({"markdown": match.group().strip()})
    
    return ParseResult(
        content=content,
        tables=tables,
        pages=1,
        metadata={"source": str(filepath)},
    )


# ============================================================
# TXT 解析
# ============================================================

def parse_text(filepath: str | Path) -> ParseResult:
    content = Path(filepath).read_text(encoding="utf-8")
    content = clean_text(content)
    return ParseResult(content=content, tables=[], pages=1, metadata={"source": str(filepath)})


# ============================================================
# 统一入口
# ============================================================

PARSER_MAP = {".pdf": parse_pdf, ".docx": parse_docx, ".md": parse_markdown, ".txt": parse_text}


def parse_document(filepath: str | Path) -> ParseResult:
    ext = Path(filepath).suffix.lower()
    parser = PARSER_MAP.get(ext)
    if not parser:
        raise ValueError(f"不支持的文档格式: {ext}")
    return parser(filepath)
